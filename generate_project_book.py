"""
generate_project_book.py  —  תיק פרויקט Super Mario  (styled DOCX)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "תיק_פרויקט_סופר_מריו.docx"

# ── colour palette ────────────────────────────────────────────────────────────
C_DARK_BLUE  = "1F3864"
C_MED_BLUE   = "2E75B6"
C_LIGHT_BLUE = "DEEAF1"
C_RED        = "C0392B"
C_ORANGE     = "E67E22"
C_AMBER_BG   = "FEF9E7"
C_GREY_DARK  = "404040"
C_GREY_MED   = "AAAAAA"
C_GREY_BG    = "F5F5F5"
C_DIAG_BG    = "EBF5FB"
C_DIAG_BD    = "2980B9"
C_WHITE      = "FFFFFF"

FONT   = "David"
BODY_SIZE = 12

# ═══════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL XML HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _rgb(hex6):
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return RGBColor(r, g, b)


def _set_rtl(para, align=WD_ALIGN_PARAGRAPH.RIGHT):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:bidi")):
        pPr.remove(old)
    pPr.append(OxmlElement("w:bidi"))
    para.alignment = align


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _para_shading(para, fill, color="auto"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), color)
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _para_border(para, sides, val="single", sz="6", color="000000", space="4"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    bd = OxmlElement("w:pBdr")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), space)
        b.set(qn("w:color"), color)
        bd.append(b)
    pPr.append(bd)


def _para_indent(para, left=0, right=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    if left:
        ind.set(qn("w:left"), str(left))
    if right:
        ind.set(qn("w:right"), str(right))
    pPr.append(ind)


def _rtl_run(para, text, bold=False, italic=False, size=BODY_SIZE,
             font=FONT, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = _rgb(color) if isinstance(color, str) else color
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), font)
    return run


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ═══════════════════════════════════════════════════════════════════════════════
#  STYLED BUILDING BLOCKS
# ═══════════════════════════════════════════════════════════════════════════════

def H1(doc, text):
    """Dark-blue 18 pt heading with a red bottom rule."""
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=240, after=80)
    _para_border(p, ["bottom"], val="single", sz="12", color=C_RED, space="4")
    _rtl_run(p, text, bold=True, size=18, color=C_DARK_BLUE)
    return p


def H2(doc, text):
    """Medium-blue 14 pt with a right-side accent bar + light-blue background."""
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=160, after=60)
    _para_shading(p, fill=C_LIGHT_BLUE)
    _para_border(p, ["right"], val="single", sz="20", color=C_MED_BLUE, space="6")
    _para_indent(p, right=160)
    _rtl_run(p, text, bold=True, size=14, color=C_MED_BLUE)
    return p


def H3(doc, text):
    """12 pt bold, dark grey."""
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=120, after=40)
    _rtl_run(p, text, bold=True, size=12, color=C_GREY_DARK)
    return p


def para(doc, text, bold=False, italic=False, size=BODY_SIZE, color=None):
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=0, after=60, line=276)   # ≈ 1.15 line spacing
    if text:
        _rtl_run(p, text, bold=bold, italic=italic, size=size,
                 color=color or C_GREY_DARK)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=0, after=40, line=276)
    _para_indent(p, right=360)
    _rtl_run(p, "• ", bold=False, size=BODY_SIZE, color=C_MED_BLUE)
    _rtl_run(p, text, size=BODY_SIZE, color=C_GREY_DARK)
    return p


def placeholder(doc, text):
    """Amber-bordered box for content the student must fill in."""
    p = doc.add_paragraph()
    _set_rtl(p)
    _para_spacing(p, before=80, after=80)
    _para_shading(p, fill=C_AMBER_BG)
    _para_border(p, ["top", "left", "bottom", "right"],
                 val="single", sz="6", color=C_ORANGE, space="4")
    _para_indent(p, right=200, left=200)
    _rtl_run(p, "✏  ", bold=True, size=11, color=C_ORANGE)
    _rtl_run(p, text, italic=True, size=11, color=C_ORANGE)
    return p


def screenshot_box(doc, description):
    """Grey box = where the student will paste a screenshot."""
    p = doc.add_paragraph()
    _set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    _para_spacing(p, before=120, after=120)
    _para_shading(p, fill=C_GREY_BG)
    _para_border(p, ["top", "left", "bottom", "right"],
                 val="single", sz="6", color=C_GREY_MED, space="4")
    _para_indent(p, right=360, left=360)
    _rtl_run(p, f"📷  [ צילום מסך: {description} ]",
             italic=True, size=11, color="777777")
    doc.add_paragraph()
    return p


def diagram_box(doc, description):
    """Blue dashed box = where the student will paste a diagram."""
    p = doc.add_paragraph()
    _set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    _para_spacing(p, before=120, after=120)
    _para_shading(p, fill=C_DIAG_BG)
    _para_border(p, ["top", "left", "bottom", "right"],
                 val="dashed", sz="6", color=C_DIAG_BD, space="4")
    _para_indent(p, right=360, left=360)
    _rtl_run(p, f"📊  [ תרשים: {description} ]",
             italic=True, size=11, color=C_DIAG_BD)
    doc.add_paragraph()
    return p


def code_line(doc, text):
    """Monospace line for code / tree structures."""
    p = doc.add_paragraph()
    _set_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
    _para_spacing(p, before=0, after=0)
    _para_shading(p, fill="F8F8F8")
    _para_indent(p, right=200, left=200)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = _rgb("2C3E50")
    return p


# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE SETUP  &  HEADER / FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

def setup_page(doc):
    s = doc.sections[0]
    s.page_height = Cm(29.7)
    s.page_width  = Cm(21)
    s.left_margin  = Cm(2.5)
    s.right_margin = Cm(2.5)
    s.top_margin   = Cm(2.5)
    s.bottom_margin = Cm(2.0)


def setup_header_footer(doc):
    section = doc.sections[0]

    # ── Header ──────────────────────────────────────────────────────────
    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_border(hp, ["bottom"], val="single", sz="4", color=C_MED_BLUE, space="4")
    run = hp.add_run("[שם התלמיד]  |  פיתוח משחק פלטפורמה בסגנון Super Mario")
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb("555555")
    run.font.name = FONT

    # ── Footer ──────────────────────────────────────────────────────────
    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(fp, ["top"], val="single", sz="4", color=C_MED_BLUE, space="4")
    run = fp.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb("555555")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def _banner(doc, text, fill, text_color=C_WHITE, size=22, before=0, after=0):
    """Full-width coloured banner paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_shading(p, fill=fill)
    _para_spacing(p, before=before, after=after)
    _rtl_run(p, text, bold=True, size=size, color=text_color)
    return p


def cover_page(doc):
    # ── top blue banner ──────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    _banner(doc, "[ לוגו בית הספר ]", fill=C_LIGHT_BLUE,
            text_color=C_DARK_BLUE, size=13, before=120, after=120)
    _banner(doc, "[ שם בית הספר ]", fill=C_DARK_BLUE,
            text_color=C_WHITE, size=18, before=0, after=0)

    for _ in range(2):
        doc.add_paragraph()

    # ── project title ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=0, after=80)
    _rtl_run(p, "תיק פרויקט בהנדסת תוכנה", bold=True, size=14, color=C_GREY_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_shading(p, fill=C_LIGHT_BLUE)
    _para_spacing(p, before=80, after=80)
    _para_border(p, ["top", "left", "bottom", "right"],
                 val="single", sz="10", color=C_DARK_BLUE, space="6")
    _rtl_run(p, "פיתוח משחק פלטפורמה בסגנון Super Mario",
             bold=True, size=20, color=C_DARK_BLUE)

    for _ in range(3):
        doc.add_paragraph()

    # ── details box ─────────────────────────────────────────────────────
    details = [
        ("שם התלמיד",   "[שם התלמיד]",          True),
        ("ת.ז.",        "[ת.ז. התלמיד]",         True),
        ("שם המנחה",    "[שם המנחה]",             True),
        ("שם החלופה",
         "שירותי אינטרנט, תכנות אסינכרוני ומסדי נתונים — הנדסת תוכנה 883589",
         False),
        ("תאריך ההגשה", "[תאריך ההגשה]",         True),
    ]
    for label, value, is_ph in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p, before=40, after=40)
        _rtl_run(p, f"{label}: ", bold=True, size=12, color=C_DARK_BLUE)
        _rtl_run(p, value, italic=is_ph, size=12,
                 color=C_ORANGE if is_ph else C_GREY_DARK)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def toc_section(doc):
    H1(doc, "תוכן עניינים")
    placeholder(doc,
        "יש ליצור תוכן עניינים אוטומטי בוורד לאחר השלמת המסמך.\n"
        "References ← Table of Contents ← Automatic Table")
    para(doc, "")
    placeholder(doc, "[ כאן יופיע תוכן העניינים האוטומטי ]")
    page_break(doc)


def intro_section(doc):
    H1(doc, "מבוא")

    H2(doc, "הרקע לפרויקט")
    para(doc, "שם הפרויקט: SuperMario — משחק פלטפורמה")
    para(doc,
        "תיאור קצר: פרויקט זה הוא משחק מחשב דו-ממדי בסגנון Nintendo Super Mario הקלאסי, "
        "הבנוי מאפס בשפת #C עם מסגרת WinForms. המשחק כולל 5 רמות (3 ידניות ו-2 "
        "נוצרות אלגוריתמית), 6 סוגי אויבים עם התנהגויות שונות, מנוע פיזיקה, "
        "מערכת ספרייטים, ומצב אימון AI מבוסס neuroevolution.")
    para(doc, "קהל היעד: חובבי משחקי וידאו ומשחקי פלטפורמה קלאסיים, גילאים 10 ומעלה.")
    placeholder(doc, "הסיבות לבחירת הנושא: [יש להשלים בנוסח אישי]")

    H2(doc, "מטרות המערכת")
    H3(doc, "מטרות על")
    bullet(doc, "פיתוח משחק פלטפורמה מלא, פונקציונלי וניתן לשחק ב-Windows")
    bullet(doc, "מימוש מודול אימון AI המסוגל ללמוד לשחק במשחק באופן אוטונומי")

    H3(doc, "מטרות נלוות")
    bullet(doc, "יישום עקרונות תכנות מונחה עצמים: ירושה, פולימורפיזם, הפשטה")
    bullet(doc, "פיתוח מנוע פיזיקה (כבידה, קפיצה, זיהוי התנגשויות)")
    bullet(doc, "ניהול מצב משחק, אנימציות, וממשק משתמש גרפי")
    bullet(doc, "מימוש אלגוריתם neuroevolution (רשת נוירונים + אלגוריתם גנטי)")

    H2(doc, "תיאור המערכת")
    para(doc,
        "המערכת בנויה כיישום שולחני ל-Windows הכולל שתי מצבי פעולה: מצב משחק רגיל "
        "ומצב אימון AI. המשחק מחולק ל-5 שכבות לוגיות: Core (ליבת המשחק), "
        "Enemies (אויבים), World (עצמי עולם), UI (ממשק משתמש), ו-ML (בינה מלאכותית).")

    H2(doc, "גבולות המערכת")
    bullet(doc, "פלטפורמה: Windows Desktop בלבד (.NET Framework 4.7.2, WinForms)")
    bullet(doc, "אין תמיכה ב-Linux / macOS")
    bullet(doc, "אין תמיכה מרובת שחקנים")
    bullet(doc, "אין שמירת התקדמות משחק")

    H2(doc, "סביבת פיתוח")
    bullet(doc, "סביבת פיתוח: Microsoft Visual Studio 2022")
    bullet(doc, ".NET Framework 4.7.2")
    bullet(doc, "מערכת גרסאות: Git / GitHub")

    H2(doc, "שפת תכנות")
    para(doc,
        "#C — שפה מונחית עצמים הנתמכת מלאית על ידי .NET Framework ו-WinForms.")

    H2(doc, "שכבות")
    for layer, desc in [
        ("Core",    "מחלקות ליבה: שחקן, אנימציה, נתוני משחק, מנהל משחק"),
        ("Enemies", "היררכיית אויבים עם ירושה"),
        ("World",   "פלטפורמות, צינורות וקצה הרמה"),
        ("UI",      "חלונות, לולאת משחק, קלט, HUD"),
        ("ML",      "רשת נוירונים, אלגוריתם גנטי, סוכן AI"),
    ]:
        bullet(doc, f"{layer}: {desc}")

    H2(doc, "פלטפורמות הלקוח")
    para(doc,
        "יישום Windows Desktop (WinForms) — נבחר בשל תמיכה מלאה ב-.NET Framework 4.7.2, "
        "כלי GUI מובנים, ויכולת ציור גרפי מתקדמת.")

    H2(doc, "אתגרים מרכזיים")
    bullet(doc,
        "פיתוח מנוע פיזיקה: מימוש כבידה, קפיצה, ואלגוריתם זיהוי התנגשויות "
        "הנכון גם עבור תנועה מהירה")
    bullet(doc,
        "ניהול 6 סוגי אויבים שונים עם התנהגויות מגוונות — "
        "פתרון: היררכיית ירושה עם Enemy כמחלקת בסיס מופשטת")
    bullet(doc,
        "מערכת ספרייטים: יצירה, טעינה ואנימציה של תמונות לכל אובייקט")
    bullet(doc,
        "מימוש Neuroevolution: שילוב רשת נוירונים עם אלגוריתם גנטי לאימון "
        "סוכן AI אוטונומי בזמן-אמת")
    placeholder(doc, "אתגרים אישיים נוספים: [יש להשלים]")

    H2(doc, "חידושים ועדכונים")
    bullet(doc, "כל ספרייט עוצב כפיקסל-ארט מקורי בסקריפט Python")
    bullet(doc, "ספרייטים נטענים דינמית ומוצגים דרך PictureBox — ללא ציור GDI+ ידני")
    bullet(doc, "רמות 4–5 נבנות אלגוריתמית מ-templates מוגדרים מראש")

    page_break(doc)


def system_analysis_section(doc):
    H1(doc, "ניתוח מערכת")

    H2(doc, "מסמך ייזום")
    placeholder(doc,
        "[יש להשלים את מסמך הייזום: שם הפרויקט, מזמין העבודה, תיאור הצורך]")
    diagram_box(doc, "מסמך ייזום — להשלים")

    H2(doc, "מצב קיים")
    para(doc,
        "לפני פיתוח הפרויקט לא קיים כלי דומה בשימוש. "
        "המשחק נבנה מאפס ואינו מחליף מערכת קיימת.")

    H2(doc, "מטרות המערכת העתידית")
    placeholder(doc, "[יש להשלים: מה המערכת אמורה לספק לאחר הפיתוח]")

    H2(doc, "מצב עתידי")
    placeholder(doc, "[יש להשלים: תיאור המצב הרצוי לאחר יישום הפרויקט]")

    H2(doc, "עץ תהליכים")
    for line in [
        "Main Menu",
        "  ├── משחק רגיל → בחירת רמה (1–5) → משחק → {ניצחון / מוות / רמה הבאה}",
        "  ├── אימון AI  → הגדרת פרמטרים → הרצת סוכנים → אבולוציה דורית",
        "  └── יציאה",
    ]:
        code_line(doc, line)
    diagram_box(doc, "עץ תהליכים — ניתן לייצא כתרשים ויזואלי ולהדביק כאן")

    H2(doc, "Use Case — תהליכי המערכת")
    H3(doc, "שחקן רגיל")
    bullet(doc, "מפעיל משחק → בוחר רמה → שולט בדמות → אוסף מטבעות ופטריות")
    bullet(doc, "קופץ על אויבים / נפגע מאויבים → מאבד חיים")
    bullet(doc, "מגיע לקצה הרמה → עובר לרמה הבאה")
    H3(doc, "מאמן AI")
    bullet(doc, "פותח מסך אימון → מגדיר פרמטרים (גודל אוכלוסייה, קצב מוטציה)")
    bullet(doc, "מריץ / מפסיק אימון → צופה בהתקדמות הסוכנים")
    bullet(doc, "שומר את הרשת הטובה ביותר לקובץ")
    diagram_box(doc, "תרשים Use Case — יש לשרטט ולהדביק כאן")

    H2(doc, "DFD-0 — תרשים זרימת נתונים")
    diagram_box(doc, "תרשים DFD-0 — יש לשרטט ולהדביק כאן")

    para(doc,
        "הערה: ERD אינו רלוונטי לפרויקט זה — המשחק אינו משתמש במסד נתונים רלציוני. "
        "כל הנתונים מנוהלים בזיכרון. ראה פרק בסיס נתונים.")

    page_break(doc)


def database_section(doc):
    H1(doc, "בסיס נתונים")

    para(doc,
        "פרויקט זה אינו עושה שימוש במסד נתונים רלציוני חיצוני. "
        "כל נתוני המשחק מנוהלים בזיכרון (in-memory) במהלך ריצת התוכנה.")

    H2(doc, "מבני נתונים עיקריים בזיכרון")
    for name, desc in [
        ("Player",         "מצב השחקן: מיקום, מהירות, בריאות, ניקוד"),
        ("Enemy (abstract)","מצב אויב בסיסי: מיקום, כיוון, מהירות, חי/מת"),
        ("Coin",           "מטבע: מיקום, האם נאסף"),
        ("Mushroom",       "פטרייה: מיקום, האם פעילה, כיוון תנועה"),
        ("QuestionBlock",  "בלוק שאלה: מיקום, האם נפתח"),
        ("PlatformData",   "נתוני פלטפורמה: X, Y, רוחב, גובה"),
    ]:
        bullet(doc, f"{name} — {desc}")

    H2(doc, "שמירת נתונים לקובץ — NetworkSerializer")
    para(doc,
        "הנתון היחיד הנשמר לדיסק הוא מצב הרשת הנוירונית של הסוכן הטוב ביותר, "
        "באמצעות המחלקה NetworkSerializer.")

    H3(doc, "פורמט קובץ הרשת הנוירונית")
    for line in [
        "SMNET1           ← סוג הקובץ (header)",
        "generation=42    ← מספר הדור",
        "fitness=1850     ← ביצועי הסוכן הטוב",
        "shape=5,8,8,3    ← צורת הרשת (שכבות ונוירונים)",
        "<layerIdx>,<neuronIdx>,<bias>,<w0>,<w1>,...  ← נתוני כל נוירון",
    ]:
        code_line(doc, line)

    para(doc,
        "מבנה זה מאפשר שמירה וטעינה מהירה של הרשת הנוירונית בין הפעלות התוכנה, "
        "ללא תלות בספריות חיצוניות.")

    page_break(doc)


def implementation_section(doc):
    H1(doc, "מימוש הפרויקט")

    # ── Core ──────────────────────────────────────────────────────────────
    H2(doc, "שכבת Core — ליבת המשחק")
    diagram_box(doc, "תרשים UML של מחלקות Core")

    H3(doc, "Player")
    para(doc,
        "מחלקת Player מנהלת את כל מצב השחקן ופיזיקתו. "
        "הפיזיקה פשוטה ובסיסית: תנועה אופקית במהירות קבועה, "
        "כבידה קבועה לכל פריים, וקפיצה המציבה מהירות אנכית שלילית.")
    bullet(doc, "מהירות הליכה: 5 פיקסלים/פריים")
    bullet(doc, "כבידה: 0.6 פיקסלים/פריים²")
    bullet(doc, "עוצמת קפיצה: -13 (כלפי מעלה)")
    bullet(doc, "מהירות נפילה מרבית: 15 פיקסלים/פריים")
    bullet(doc, "מצב: בריאות (3 לבבות), ניקוד, IsGrounded")

    H3(doc, "Animator")
    para(doc,
        "מחלקת Animator מנהלת מחזור פריימים לאנימציה. "
        "מקבלת מערך של תמונות ומחזירה את התמונה הנוכחית לפי קצב מוגדר.")

    H3(doc, "GameData")
    para(doc,
        "מגדירה מבני נתונים לפריטי המשחק: Coin, Mushroom, QuestionBlock. "
        "כל אחד מכיל מיקום, גודל ומצב.")

    H3(doc, "Sprites")
    para(doc,
        "טוען ומאחסן את כל תמונות המשחק (PNG) פעם אחת בהפעלה. "
        "כל האובייקטים מקבלים את תמונתם מאוסף זה — "
        "אין ציור GDI+ ידני בשכבת המשחק.")

    # ── Enemies ──────────────────────────────────────────────────────────
    H2(doc, "שכבת Enemies — מערכת האויבים")
    para(doc,
        "מערכת האויבים בנויה על היררכיית ירושה. "
        "כל ההתנהגות המשותפת מרוכזת במחלקות הבסיס, "
        "ותת-מחלקה מגדירה רק מהירות, ספרייטים והתנהגות מיוחדת.")
    diagram_box(doc, "תרשים ירושה — Enemy hierarchy")

    for line in [
        "Enemy (abstract)  ←  מחלקת בסיס לכל האויבים",
        "  ├── SquishableEnemy (abstract)  ←  אויבים הניתנים למעיכה",
        "  │     ├── Goomba         — פטרייה, הולכת לאט",
        "  │     ├── FastEnemy      — חיפושית אדומה, מהירה פי 2",
        "  │     ├── JumpingEnemy   — קופצת מדי 2 שניות, כחולה",
        "  │     └── PlatformPatrolEnemy — עוצרת בקצה פלטפורמה",
        "  ├── Koopa          — צב בקליפה",
        "  └── FlyingEnemy    — קואפה מעופפת, 2 מכות להריגה",
    ]:
        code_line(doc, line)

    para(doc, "")
    for name, desc in [
        ("Goomba",               "אויב בסיסי — הולך שמאלה, מת ממעיכה"),
        ("FastEnemy",            "פי 2 מהיר מגומבה — עיצוב חיפושית אדומה"),
        ("JumpingEnemy",         "קופץ אוטומטית כל 2 שניות — עיצוב כחול עם רגלי קפיץ"),
        ("PlatformPatrolEnemy",  "מזהה קצה פלטפורמה — מתהפך לפני נפילה"),
        ("Koopa",                "מת ממכה ראשונה ועובר למצב קליפה גולשת"),
        ("FlyingEnemy",          "טיסה בדפוס סינוס, דורש 2 מכות להריגה"),
    ]:
        bullet(doc, f"{name}: {desc}")

    # ── World ─────────────────────────────────────────────────────────────
    H2(doc, "שכבת World")
    para(doc,
        'המחלקה GameObjectS עוטפת כל פלטפורמה או צינור בעולם המשחק. '
        'היא מכילה PictureBox עם BackgroundImage משובץ (tile), ותחום (Type) שיכול '
        'להיות "ground", "pipe", או "finish".')

    # ── UI ────────────────────────────────────────────────────────────────
    H2(doc, "שכבת UI — ממשק המשתמש ולולאת המשחק")
    diagram_box(doc, "תרשים זרימה בין מסכי המשחק")

    H3(doc, "MainMenuForm — מסך פתיחה")
    para(doc,
        "מסך הפתיחה מציג 4 כפתורים: משחק רגיל, אימון AI, עזרה, ויציאה. "
        "מפעיל את mainWin או TrainingForm בהתאם לבחירה.")

    H3(doc, "mainWin — חלון המשחק הראשי (8 קבצים partial)")
    para(doc, "לולאת המשחק פועלת על Timer בתדר 16ms (≈60fps):")
    for line in [
        "Timer.Tick → ReadInput() → Player.Move() → UpdateEnemies()",
        "          → CheckCollisions() → UpdateAnimatedSprites()",
        "          → ScrollCamera() → UpdateHUD()",
    ]:
        code_line(doc, line)

    para(doc, "")
    bullet(doc, "mainWin.cs — שדות, לולאה, קלט")
    bullet(doc, "mainWin.Physics.cs — זיהוי התנגשויות, מצלמה, מוות")
    bullet(doc, "mainWin.EnemyUpdates.cs — עדכון כל סוגי האויבים")
    bullet(doc, "mainWin.EnemyPhysics.cs — פיזיקת אויבים")
    bullet(doc, "mainWin.Collectibles.cs — איסוף מטבעות ופטריות")
    bullet(doc, "mainWin.HUD.cs — HUD: לבבות, ניקוד, ספרייטים")
    bullet(doc, "mainWin.LevelBuilder.cs — בניית / ניקוי רמה")
    bullet(doc, "mainWin.LevelData.cs — נתוני הרמות הסטטיים")

    H3(doc, "TrainingForm — מסך אימון AI")
    para(doc,
        "מסך מפוצל: קנבס שמאל לסימולציה חיה של עד 60 סוכנים, "
        "ולוח בקרה ימני עם: מספר דור, כמה חיים, ביצועי ניצחון, "
        "הגדרות פרמטרים, ושליטה (Start/Pause/Reset).")

    # ── ML ────────────────────────────────────────────────────────────────
    H2(doc, "שכבת ML — בינה מלאכותית (Neuroevolution)")
    para(doc,
        "שכבת ה-ML מורכבת ממנוע רשת נוירונים ואלגוריתם גנטי. "
        "חשוב לציין: המחלקות Neuron, Layer, NeuralNetwork, "
        "NeuralNetworkControl, Population ו-NetParams "
        "הותאמו מפרויקט Flappy Bird קיים (ראה ביבליוגרפיה).")
    diagram_box(doc, "ארכיטקטורת מודול ה-ML")

    H3(doc, "MarioAgent — מחלקה מקורית")
    para(doc,
        "MarioAgent היא המחלקה המקורית שנכתבה לפרויקט זה. "
        "כל סוכן הוא לואיג'י עצמאי עם פיזיקה ורשת נוירונים משלו.")
    bullet(doc, "קלטים (5): מרחק לאויב הקרוב, גובה השחקן, מהירות, האם על הקרקע, מרחק לבור")
    bullet(doc, "פלטים (3): קפיצה, תנועה שמאלה, תנועה ימינה")
    bullet(doc, "פיטנס: מרחק X שהגיע + 50 נקודות לכל מטבע שנאסף")

    H3(doc, "Population — אלגוריתם גנטי (מותאם)")
    para(doc,
        "מנהל את כלל הסוכנים: בוחר את הטובים ביותר, "
        "מבצע crossover בין רשתות, מוסיף מוטציות אקראיות.")

    page_break(doc)


def user_guide_section(doc):
    H1(doc, "מדריך למשתמש")

    para(doc,
        "המערכת תומכת בשני סוגי משתמשים: שחקן רגיל ומאמן AI. "
        "לכל סוג ממשק מותאם עם אפשרויות שונות.")
    para(doc, "הפרויקט נבדק על: Windows 10 / Windows 11 (64-bit), .NET Framework 4.7.2 ומעלה.")

    H2(doc, "שחקן רגיל")

    H3(doc, "מסך ראשי")
    screenshot_box(doc, "מסך Main Menu — 4 כפתורי ניווט")
    para(doc, "מהמסך הראשי ניתן:")
    bullet(doc, "לחיצה על 'שחק' — כניסה למשחק")
    bullet(doc, "לחיצה על 'אימון AI' — מעבר למסך אימון הסוכן")
    bullet(doc, "לחיצה על 'עזרה' — הצגת הוראות")
    bullet(doc, "לחיצה על 'יציאה' — סגירת התוכנה")

    H3(doc, "מסך המשחק")
    screenshot_box(doc, "מסך משחק — שחקן, אויבים, פלטפורמות ו-HUD עם לבבות וניקוד")
    para(doc, "שליטה:")
    bullet(doc, "חץ שמאל / ימין — תנועה")
    bullet(doc, "רווח או חץ למעלה — קפיצה (רק כשעל הקרקע)")
    bullet(doc, "ESC — יציאה מהמשחק")
    para(doc, "HUD (ראש הדף):")
    bullet(doc, "3 לבבות — מייצגים את החיים שנותרו")
    bullet(doc, "ניקוד — מתעדכן בכל מטבע שנאסף")
    bullet(doc, "מספר רמה נוכחית")
    para(doc,
        "אם השחקן נגמרים לו החיים — מוצגת הודעה ומתחיל משחק חדש. "
        "אם מגיע לדגל הסיום — עובר לרמה הבאה.")

    H3(doc, "מסך סיום")
    screenshot_box(doc, "מסך סיום רמה או מוות — הודעה וחזרה לתפריט")

    H2(doc, "מאמן AI")

    H3(doc, "מסך אימון AI")
    screenshot_box(doc, "מסך אימון — קנבס סימולציה שמאל, לוח בקרה ורשת נוירונים ימין")
    para(doc, "פרמטרים ניתנים להגדרה:")
    bullet(doc, "גודל אוכלוסייה — כמה סוכנים בכל דור")
    bullet(doc, "אחוז מוטציה — כמה שינוי אקראי בין דורות")
    bullet(doc, "אחוז שרידים — כמה מהטובים עוברים לדור הבא")
    bullet(doc, "צורת הרשת — מספר הנוירונים בכל שכבה")
    para(doc, "כפתורי שליטה:")
    bullet(doc, "Start / Pause — הפעלה / השהייה")
    bullet(doc, "Reset — איפוס האימון")
    bullet(doc, "Apply Settings — החלת הגדרות חדשות")
    bullet(doc, "שמירת הרשת הטובה לקובץ .smnet")

    page_break(doc)


def developer_guide_section(doc):
    H1(doc, "מדריך למפתח")

    H2(doc, "הגדרת סביבת הפיתוח")
    bullet(doc, "דרישות: Windows 10/11, Visual Studio 2022, .NET Framework 4.7.2, Python 3.x")
    bullet(doc, "פתיחת הפרויקט: supermario.sln → F5 להרצה")
    bullet(doc,
        "יצירת ספרייטים: python generate_spritesheets.py "
        "→ PNG נוצרים ב-assets/textures/sprites/")

    H2(doc, "מסך ראשי — MainMenuForm")
    screenshot_box(doc, "Main Menu — מנקודת מבט מפתח")
    para(doc, "קובץ: UI/MainMenuForm.cs  |  מחלקה: MainMenuForm")
    bullet(doc, "4 כפתורים: btnPlay, btnTrain, btnHelp, btnExit")
    bullet(doc, "btnPlay: new mainWin().Show(), סגירת MainMenuForm")
    bullet(doc, "btnTrain: new TrainingForm().Show(), סגירת MainMenuForm")
    bullet(doc, "הוספת כפתור: גרור Button ב-Designer, הוסף EventHandler")

    H2(doc, "מסך המשחק — mainWin")
    screenshot_box(doc, "מסך משחק — מנקודת מבט מפתח: PictureBox של שחקן ואויבים")
    para(doc,
        "mainWin היא partial class המחולקת ל-8 קבצים. "
        "הרכיב המרכזי הוא gameTimer (Timer, 16ms) המפעיל את כל הלוגיקה.")
    para(doc, "כיצד להוסיף רמה חדשה:")
    bullet(doc, "פתח mainWin.LevelData.cs והוסף מערך PlatformData[] חדש")
    bullet(doc, "הוסף case ב-switch של BuildLevel() ב-mainWin.LevelBuilder.cs")

    H2(doc, "מערכת האויבים")
    screenshot_box(doc, "6 סוגי האויבים בפעולה")
    para(doc, "כיצד להוסיף סוג אויב חדש:")
    bullet(doc,
        "צור מחלקה חדשה היורשת מ-SquishableEnemy (אם ניתן למעיכה) "
        "או מ-Enemy (אחרת)")
    bullet(doc, "הגדר מהירות, ספרייטים וגודל בבנאי")
    bullet(doc, "דרוס את Update() להתנהגות מיוחדת (קפיצה, טיסה)")
    bullet(doc, "הוסף List<NewEnemy> ב-mainWin.cs")
    bullet(doc, "הוסף ספאווון ב-LevelBuilder.cs ועדכון ב-EnemyUpdates.cs")
    para(doc, "כיצד לשנות ספרייט:")
    bullet(doc, "ערוך את generate_spritesheets.py והרץ: python generate_spritesheets.py")

    H2(doc, "מסך אימון AI — TrainingForm + ML")
    screenshot_box(doc, "מסך אימון — תצוגת הרשת הנוירונית ולוח הבקרה")
    para(doc,
        "TrainingForm מנהלת את הציור (RenderScene) ואת ה-Population. "
        "MarioAgent מגדיר את קלטי הרשת והפיזיקה של הסוכן.")
    para(doc, "כיצד לשנות את מספר הקלטים:")
    bullet(doc, "ערוך MarioAgent.GetInputs() ב-ML/MarioAgent.cs")
    bullet(doc, "עדכן NetParams.DefaultShape[0] ב-ML/NetParams.cs")
    bullet(doc, "איפוס האימון נדרש לאחר שינוי")

    H2(doc, "מחלקות שנבנו על ידי המפתח")
    for cls in [
        "Core/Player.cs", "Core/Animator.cs", "Core/GameData.cs",
        "Core/GameManager.cs", "Core/Sprites.cs",
        "Enemies/Enemy.cs", "Enemies/SquishableEnemy.cs",
        "Enemies/Goomba.cs", "Enemies/Koopa.cs", "Enemies/FastEnemy.cs",
        "Enemies/JumpingEnemy.cs", "Enemies/PlatformPatrolEnemy.cs",
        "Enemies/FlyingEnemy.cs",
        "World/GameObjectS.cs",
        "UI/mainWin.cs (+ 7 partial files)", "UI/MainMenuForm.cs",
        "UI/TrainingForm.cs",
        "ML/MarioAgent.cs",
    ]:
        bullet(doc, cls)

    page_break(doc)


def reflection_section(doc):
    H1(doc, "סיכום אישי / רפלקציה")
    para(doc,
        "חלק זה מיועד לתיאור אישי של תהליך הפיתוח. "
        "יש למלא לפחות חצי עמוד עד עמוד שלם.")

    H2(doc, "תיאור תהליך העבודה")
    placeholder(doc,
        "[תאר את תהליך העבודה על הפרויקט: מה עשית ראשון, "
        "כיצד התקדמת, מה שינית בדרך]")

    H2(doc, "אתגרים ודרכי פתרון")
    placeholder(doc,
        "[תאר את האתגרים המרכזיים שנתקלת בהם ואיך פתרת אותם. "
        "לדוגמה: בעיות בפיזיקה, באויבים, בביצועים]")

    H2(doc, "תהליך למידה")
    placeholder(doc,
        "[מה למדת באופן עצמאי שלא נלמד בכיתה? "
        "אילו מושגים, טכנולוגיות או כלים גילית בכוחות עצמך?]")

    H2(doc, "כלים שאקח להמשך")
    placeholder(doc,
        "[אילו כלים, עקרונות או ידע שרכשת תשתמש בהם בפרויקטים הבאים?]")

    H2(doc, "בראייה לאחור")
    placeholder(doc,
        "[אם היית מתחיל מחדש, מה היית עושה אחרת? "
        "אילו החלטות עיצוביות היית משנה?]")

    H2(doc, "שיפורים אפשריים")
    placeholder(doc,
        "[אם היו לך יותר זמן ומשאבים, כיצד היית משפר את הפרויקט? "
        "לדוגמה: רמות נוספות, סוגי אויבים, מצב multiplayer]")

    page_break(doc)


def bibliography_section(doc):
    H1(doc, "ביבליוגרפיה")
    para(doc, "מקורות המידע שנעשה בהם שימוש בפרויקט, לפי כללי APA:")

    bullet(doc,
        "[מחבר פרויקט Flappy Bird]. ([שנה]). [שם פרויקט Flappy Bird עם neuroevolution]. "
        "GitHub. [יש להשלים את הקישור המדויק]")
    bullet(doc,
        "Microsoft. (2024). Windows Forms documentation. Microsoft Docs. "
        "https://docs.microsoft.com/en-us/dotnet/desktop/winforms/")
    bullet(doc,
        "Microsoft. (2024). C# language reference. Microsoft Docs. "
        "https://docs.microsoft.com/en-us/dotnet/csharp/")
    bullet(doc,
        "Nintendo. (1985). Super Mario Bros. [Video game]. Nintendo.")
    placeholder(doc, "[יש להוסיף מקורות נוספים שהשתמשת בהם]")

    page_break(doc)


def appendix_section(doc):
    H1(doc, "נספחים")

    H2(doc, "נספח א׳ — קוד מחלקות הפרויקט")
    para(doc,
        "יש לצרף כאן את קוד כל מחלקות הפרויקט. "
        "הקוד יוצג כתדפיס מסודר (לא צילום מסך) עם הערות בכל המקומות הרלוונטיים.")
    for cls in [
        "Core/Player.cs", "Core/Animator.cs", "Core/GameData.cs",
        "Core/GameManager.cs", "Core/Sprites.cs",
        "Enemies/Enemy.cs", "Enemies/SquishableEnemy.cs",
        "Enemies/Goomba.cs", "Enemies/Koopa.cs", "Enemies/FastEnemy.cs",
        "Enemies/JumpingEnemy.cs", "Enemies/PlatformPatrolEnemy.cs",
        "Enemies/FlyingEnemy.cs",
        "World/GameObjectS.cs",
        "UI/mainWin.cs", "UI/mainWin.LevelBuilder.cs", "UI/mainWin.LevelData.cs",
        "UI/mainWin.Physics.cs", "UI/mainWin.EnemyUpdates.cs",
        "UI/mainWin.EnemyPhysics.cs", "UI/mainWin.Collectibles.cs", "UI/mainWin.HUD.cs",
        "UI/MainMenuForm.cs", "UI/TrainingForm.cs",
        "ML/MarioAgent.cs",
    ]:
        bullet(doc, cls)

    H2(doc, "נספח ב׳ — הסבר על הטכנולוגיות")
    placeholder(doc,
        "[ניתן להוסיף כאן הסברים על טכנולוגיות ששימשו בפרויקט: "
        "WinForms, GDI+, Neuroevolution, Genetic Algorithm וכו']")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()
    setup_page(doc)
    setup_header_footer(doc)

    cover_page(doc)
    toc_section(doc)
    intro_section(doc)
    system_analysis_section(doc)
    database_section(doc)
    implementation_section(doc)
    user_guide_section(doc)
    developer_guide_section(doc)
    reflection_section(doc)
    bibliography_section(doc)
    appendix_section(doc)

    doc.save(OUTPUT_PATH)
    print(f"✓ נשמר: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
